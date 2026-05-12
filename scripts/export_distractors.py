#!/usr/bin/env python3
"""
One-off script: queries the DB for every quiz question and exports them
to a Word doc with the correct answer + 3 distractors clearly marked.

Run with: python3 scripts/export_distractors.py
Output:    ./distractors-review.docx

The doc is designed for Track Changes review — mark up any distractor
you want fixed and we'll batch-update from your markup.
"""

import json
import os
import re
import sys
from pathlib import Path

import psycopg2
import psycopg2.extras
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Load .env DATABASE_URL --------------------------------------------------

def load_database_url() -> str:
    env_path = Path(".env")
    if not env_path.exists():
        sys.exit("ABORT: .env not found in current directory. Run from project root.")
    for line in env_path.read_text().splitlines():
        if line.startswith("DATABASE_URL="):
            return line.split("=", 1)[1].strip().strip('"').strip("'")
    sys.exit("ABORT: DATABASE_URL not found in .env")


# --- Pull data ---------------------------------------------------------------

def fetch_data(database_url: str):
    conn = psycopg2.connect(database_url)
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

    cur.execute("SELECT id, title, discipline FROM scenarios ORDER BY title;")
    scenarios = cur.fetchall()

    cur.execute("""
        SELECT id, scenario_id, step_order, prompt, correct_actions, why_it_matters,
               questions, distractors
        FROM scenario_steps
        ORDER BY scenario_id, step_order;
    """)
    steps = cur.fetchall()
    cur.close()
    conn.close()
    return scenarios, steps


# --- Doc helpers -------------------------------------------------------------

def add_shaded_paragraph(doc, text_runs, fill_hex):
    """Adds a paragraph with a background fill color."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)
    for run_text, run_opts in text_runs:
        run = p.add_run(run_text)
        font = run.font
        font.name = "Arial"
        font.size = Pt(11)
        if "bold" in run_opts and run_opts["bold"]:
            font.bold = True
        if "color" in run_opts:
            font.color.rgb = RGBColor.from_string(run_opts["color"])
    return p


def add_left_border(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def add_indent(paragraph, twentieths_of_point=360):
    pPr = paragraph._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(twentieths_of_point))
    pPr.append(ind)


def add_styled_paragraph(doc, text_runs, indent=False, left_border_color=None, fill_hex=None):
    p = doc.add_paragraph()
    if fill_hex:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        pPr.append(shd)
    if left_border_color:
        add_left_border(p, left_border_color)
    if indent:
        add_indent(p, 360)
    for run_text, run_opts in text_runs:
        run = p.add_run(run_text)
        font = run.font
        font.name = "Arial"
        font.size = Pt(11)
        if run_opts.get("bold"):
            font.bold = True
        if run_opts.get("italic"):
            font.italic = True
        if "color" in run_opts:
            font.color.rgb = RGBColor.from_string(run_opts["color"])
    return p


# --- Main --------------------------------------------------------------------

ACCENT_HEX = "2E5BFF"
GREEN_HEX = "2E7D32"
GREEN_FILL = "E8F5E9"
GRAY_HEX = "555555"
PROMPT_FILL = "F5F5F5"


def main():
    database_url = load_database_url()
    print("Connecting to DB...")
    scenarios, steps = fetch_data(database_url)
    print(f"  {len(scenarios)} scenarios, {len(steps)} steps")

    # Group steps by scenario
    steps_by_scenario = {}
    for step in steps:
        steps_by_scenario.setdefault(step["scenario_id"], []).append(step)

    # Build entries
    entries = []
    for scenario in scenarios:
        scenario_steps = steps_by_scenario.get(scenario["id"], [])
        scenario_steps.sort(key=lambda s: s["step_order"])
        for step in scenario_steps:
            questions = step["questions"]
            if isinstance(questions, list) and len(questions) > 0:
                for idx, q in enumerate(questions):
                    if not isinstance(q, dict):
                        continue
                    q_prompt = (q.get("prompt") or "").strip()
                    q_actions = q.get("correctActions") or []
                    q_distractors = q.get("distractors") or []
                    if not q_prompt or not q_actions or len(q_distractors) != 3:
                        continue
                    entries.append({
                        "scenario_title": scenario["title"],
                        "question_label": f"Step {step['step_order']}, Q{idx + 1}",
                        "prompt": q_prompt,
                        "correct": str(q_actions[0]),
                        "distractors": [str(d) for d in q_distractors],
                    })
            else:
                s_prompt = (step["prompt"] or "").strip()
                s_actions = step["correct_actions"] or []
                s_distractors = step["distractors"] or []
                if not s_prompt or not s_actions or len(s_distractors) != 3:
                    continue
                entries.append({
                    "scenario_title": scenario["title"],
                    "question_label": f"Step {step['step_order']}",
                    "prompt": s_prompt,
                    "correct": str(s_actions[0]),
                    "distractors": [str(d) for d in s_distractors],
                })

    print(f"  {len(entries)} reviewable questions")

    if not entries:
        sys.exit("No entries to export.")

    # Build the doc
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Simtura Distractor Review")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Generated for review — {len(entries)} questions")
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor.from_string("777777")
    subtitle_run.font.size = Pt(10)

    # Instructions
    instr_head = doc.add_paragraph()
    instr_head_run = instr_head.add_run("How to use this doc")
    instr_head_run.font.size = Pt(14)
    instr_head_run.font.bold = True
    instr_head_run.font.color.rgb = RGBColor.from_string(ACCENT_HEX)

    doc.add_paragraph(
        "Turn on Track Changes (Word: Review → Track Changes → All Markup). "
        "For each question, the correct answer is highlighted green. The 3 distractors follow. "
        "Mark any distractor that is:"
    )

    issues = [
        "Too obviously wrong (gives away the answer)",
        "Too obviously right (could be marked correct by a knowledgeable person)",
        "Clinically dangerous if memorized as a 'fact' (would teach wrong information)",
        "Off-format compared to the correct answer (wildly different length/style)",
        "Repetitive across questions",
        "Outside the scope (nursing answer on EMT question, etc.)",
    ]
    for issue in issues:
        p = doc.add_paragraph(issue, style="List Bullet")

    closing = doc.add_paragraph()
    closing_run = closing.add_run(
        "Strike through or rewrite using Track Changes. Add comments for context. We'll batch-fix from your markup."
    )
    closing_run.font.italic = True
    closing_run.font.color.rgb = RGBColor.from_string(GRAY_HEX)
    closing_run.font.size = Pt(10)

    # Group entries by scenario for output
    by_scenario = {}
    for e in entries:
        by_scenario.setdefault(e["scenario_title"], []).append(e)

    question_counter = 0
    for scenario_title, scenario_entries in by_scenario.items():
        # Page break before each scenario (except first)
        if question_counter > 0:
            doc.add_page_break()

        # Scenario heading
        h = doc.add_paragraph()
        h_run = h.add_run(scenario_title)
        h_run.font.size = Pt(18)
        h_run.font.bold = True

        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"{len(scenario_entries)} questions")
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor.from_string(GRAY_HEX)
        meta_run.font.size = Pt(10)

        for e in scenario_entries:
            question_counter += 1

            # Question label
            label_p = doc.add_paragraph()
            label_p.paragraph_format.space_before = Pt(12)
            label_run = label_p.add_run(f"Q{question_counter} — ")
            label_run.font.size = Pt(11)
            label_run.font.bold = True
            label_run.font.color.rgb = RGBColor.from_string(ACCENT_HEX)
            label_run2 = label_p.add_run(e["question_label"])
            label_run2.font.size = Pt(11)
            label_run2.font.bold = True

            # Prompt (in a light gray block)
            add_styled_paragraph(
                doc,
                [(e["prompt"], {})],
                fill_hex=PROMPT_FILL,
            )

            # Correct answer (green left border + light green fill)
            add_styled_paragraph(
                doc,
                [
                    ("Correct:  ", {"bold": True, "color": GREEN_HEX}),
                    (e["correct"], {}),
                ],
                indent=True,
                left_border_color=GREEN_HEX,
                fill_hex=GREEN_FILL,
            )

            # 3 distractors
            for i, d in enumerate(e["distractors"]):
                add_styled_paragraph(
                    doc,
                    [
                        (f"Distractor {i + 1}:  ", {"bold": True, "color": GRAY_HEX}),
                        (d, {}),
                    ],
                    indent=True,
                )

    output_path = "distractors-review.docx"
    doc.save(output_path)
    print(f"\n✓ Wrote {output_path}")
    print(f"  Open in Word, turn on Track Changes, start reviewing.")


if __name__ == "__main__":
    main()
